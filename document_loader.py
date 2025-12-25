"""
文档处理模块 - 解析PDF和DOCX文件
"""
import os
from typing import List
import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import config


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "！", "？", "；", ".", "!", "?", ";", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> str:
        """加载PDF文件"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            print(f"加载PDF文件失败 {file_path}: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """加载DOCX文件"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"加载DOCX文件失败 {file_path}: {e}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """根据文件类型加载文档"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext == ".docx":
            return self.load_docx(file_path)
        else:
            print(f"不支持的文件类型: {ext}")
            return ""
    
    def load_all_documents(self, root_dir: str) -> List[LangchainDocument]:
        """递归加载目录下所有文档"""
        documents = []
        
        for root, dirs, files in os.walk(root_dir):
            for file in files:
                if file.endswith((".pdf", ".docx")):
                    file_path = os.path.join(root, file)
                    print(f"正在加载: {file_path}")
                    
                    text = self.load_document(file_path)
                    if text.strip():
                        # 创建文档对象，保存元数据
                        doc = LangchainDocument(
                            page_content=text,
                            metadata={
                                "source": file,
                                "file_path": file_path
                            }
                        )
                        documents.append(doc)
        
        return documents
    
    def split_documents(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """将文档切分成小块"""
        return self.text_splitter.split_documents(documents)


if __name__ == "__main__":
    # 测试文档加载
    loader = DocumentLoader()
    docs = loader.load_all_documents(config.DOCS_DIR)
    print(f"\n加载了 {len(docs)} 个文档")
    
    chunks = loader.split_documents(docs)
    print(f"切分成 {len(chunks)} 个文本块")
    
    if chunks:
        print(f"\n示例文本块:")
        print(f"来源: {chunks[0].metadata['source']}")
        print(f"内容: {chunks[0].page_content[:200]}...")
